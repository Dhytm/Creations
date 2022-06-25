
#https://python-docx.readthedocs.io/en/latest/
#pip intall python-docx

from turtle import width
from click import style
from docx import Document
#
from datetime import datetime
now = datetime.now()
#
from docx.shared import Pt, RGBColor, Cm # Valores de formatação
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#criação do documento ejm Python:
documento = Document()

for section_ in documento.sections:
    section_.top_margin = Cm(0.5)
    section_.bottom_margin = Cm(1)
    section_.right_margin = Cm(1)
    section_.left_margin = Cm(1)

#editar...

text_one = '''
14.04.2022 - 20:37h
Processo de criação iniciado...
Dado a ser imprimido: \n
Na linha do tempo, você se localiza em:
'''

text_tr ='''
{}
'''.format(now)

p = documento.add_paragraph(text_one)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p.add_run(f'{text_tr}').bold = True

#formatação:

p.style = documento.styles.add_style('estilo__', WD_STYLE_TYPE.PARAGRAPH)
p.style.font.name = 'Corbel Light'
p.style.font.size = Pt(12)

p.style.font.bold = False # negrito
p.style.font.italic = False # itálico

p.style.font.color.rgb = RGBColor(255, 95, 10)

# novo paragraph

for estilo in documento.styles:
    print(estilo)

p = documento.add_paragraph('PS: continuação (segundo parágrafo)', 'estilo__')

'''
template = Document('template.docx')
'''

#Inserir imagem:

p.style = documento.styles.add_style('estilo2__', WD_STYLE_TYPE.PARAGRAPH)
p.style.font.name = 'Corbel Light'
#p.style.font.size = Pt(12)

p = documento.add_paragraph('Imagem para seleção:', 'estilo2__')

documento.add_picture('imagem.png', width = Cm(4), height = Cm(4))
#Inserir tabela:

records = (
    (3, '111', 'spam'),
    (6, '121', 'coffee'),
    (9, '131', 'spam, coffee, spam and spam')
)

table = documento.add_table(rows = 1, cols = 3, style = 'Dark List Accent 3')
hdr_cells = table.rows[0].cells

hdr_cells[0].text = 'QTY'
hdr_cells[1].text = 'ID'
hdr_cells[2].text = 'DESC'

for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text =  desc

#estilo de tabela:

for estilo in documento.styles:
    print(estilo)


#Eemplementação do código em Python no Texto.docx
documento.save('creando.docx')





